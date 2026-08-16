import docx
from django.db import models
import PyPDF2

import os
import docx
import PyPDF2
from django.db import models

class Document(models.Model):
    title = models.CharField(max_length=255, verbose_name="Title")
    file = models.FileField(upload_to='documents/', verbose_name="Document File")
    content = models.TextField(blank=True, null=True, verbose_name="Extracted Content")
    uploaded_at = models.DateTimeField(auto_now_add=True, verbose_name="Uploaded At")

    def save(self, *args, **kwargs):
        super().save(*args, **kwargs)
    
        if self.file and not self.content:
            try:
                ext = os.path.splitext(self.file.name)[1].lower()
                
                if ext == '.pdf':
                    reader = PyPDF2.PdfReader(self.file.path)
                    extracted = [page.extract_text() for page in reader.pages if page.extract_text()]
                    self.content = '\n'.join(extracted)
                elif ext == '.txt':
                    with open(self.file.path, 'r', encoding='utf-8') as f:
                        self.content = f.read()
                elif ext == '.docx':
                    doc = docx.Document(self.file.path)
                    self.content = '\n'.join([para.text for para in doc.paragraphs])
                else:
                    self.content = "Unsupported format."

                super().save(update_fields=['content'])
                
                from .vector_store import process_and_store_document
                process_and_store_document(self.id, self.content, self.title)
                
            except Exception as e:
                print(f"🚨 ERROR: {e}")

    def __str__(self):
        return self.title

    class Meta:
        verbose_name = "Document"
        verbose_name_plural = "Documents"


class QAHistory(models.Model):
    question = models.TextField(verbose_name="Question")
    answer = models.TextField(verbose_name="Answer")
    created_at = models.DateTimeField(auto_now_add=True, verbose_name="Created At")

    def __str__(self):
        return f"{self.question[:50]}..."

    class Meta:
        verbose_name = "QA History"
        verbose_name_plural = "QA Histories"